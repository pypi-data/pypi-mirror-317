# coding=utf-8
"""
@File   :word_tools.py
@Time   :2024/12/28 14:12 15:20
@Author :Sunmy
@Description: word 文件处理通用工具类
"""
import string

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
# 在doc开头调用进行设置
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Mm, Inches

from .common_tools import *


def word_table_is_valid(table, value: str) -> bool:
    """
    判断表格是否有效，即表格第1行、第1列的值是否与关键字一致。
    :param table: 要检查的表格对象，例如使用 Document(file_path) 后的表格对象。
    :param value: 要检查的关键字，字符串类型。
    :return: 如果表格有效，返回 True，否则返回 False。
    """
    try:
        cell_text = common_clean_text(table.rows[0].cells[0].text)
        return cell_text == value
    except IndexError:
        return False


def word_add_custom_heading(document, level, text):
    """
    向文档添加自定义级别的标题。
    :param document: 文档对象。
    :param level: 标题级别，整数值，1 表示一级标题，2 表示二级标题，以此类推。
    :param text: 标题文本。
    """
    try:
        if level < 1 or level > 9:
            raise ValueError("Heading level must be between 1 and 9")
        else:
            heading_style = f"Heading {level}"
            document.add_heading(text, level=level).style = heading_style

    except ValueError as ve:
        print(f"ValueError in add_custom_heading: {ve}")


def word_add_custom_paragraph(document, text, style=None):
    """
    向文档中添加自定义样式的段落。
    :param document: Document对象，表示要添加段落的文档。
    :param text: 要添加的文本内容。
    :param style: 要应用的样式名称，默认为None。
    """
    try:
        if style:
            document.add_paragraph(text, style=style)
        else:
            document.add_paragraph(text)

    except Exception as e:
        print(f"An error occurred in add_custom_paragraph: {e}")


def word_add_custom_table(
        document,
        row_count,
        column_count,
        style="Table Grid",
        font_size=10.5,
        spacing_before=0,
):
    """
    向文档中添加自定义样式的表格。
    :param document: Document对象，表示要添加表格的文档。
    :param row_count: 表格的行数。
    :param column_count: 表格的列数。
    :param style: 要应用的表格样式名称，默认为 "Table Grid"。
    :param font_size: 字体大小，默认为10.5。
    :param spacing_before: 段前间距，默认为0。
    """
    try:
        table = document.add_table(rows=row_count, cols=column_count)
        if style:
            table.style = style
        else:
            table.style = "Table Grid"
        # table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)  # 设置字体大小
                    paragraph.paragraph_format.space_before = Pt(
                        spacing_before
                    )  # 设置段前间距
        return table

    except Exception as e:
        print(f"An error occurred in add_custom_table: {e}")


def word_get_section_number(paragraph_text):
    """
    从段落文本中提取章节号
    """
    try:
        parts = paragraph_text.split()
        for part in parts:
            if part.replace(".", "").isdigit():
                return part
        return None
    except Exception as e:
        print(f"An error occurred in get_section_number: {e}")
        return None


def word_add_figure_caption(document, caption_text):
    """
    在 Word 文档中添加图序号题注
    :param document: Word 文档对象
    :param caption_text: 题注文本
    """
    try:
        # 获取所有表格
        tables = document.tables

        # 获取表格数量
        table_count = len(tables)

        # 获取最近的标题（章节号）
        section_number = None
        for paragraph in document.paragraphs[::-1]:  # 倒序查找
            if paragraph.style.name.startswith("Heading"):
                section_text = word_get_section_number(paragraph.text)
                if section_text is not None:
                    section_number = section_text.split(".")[0]
                break

        # 创建一个新的段落对象用于存放图序号题注
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐

        # 创建一个新的图序号题注对象
        caption = OxmlElement("w:fldSimple")
        caption.set(qn("w:instr"), f"SEQ Figure \\* ARABIC")
        run = p.add_run()
        run.text = f"图 {section_number}-{table_count + 1}：{caption_text}"
        run.font.name = "宋体"
        run.font.size = Pt(10.5)
        p._p.append(caption)

    except Exception as e:
        print(f"An error occurred in add_figure_caption: {e}")


def word_add_indented_paragraph(document, text):
    """
    在 Word 文档中添加一个带有序号和缩进的段落。
    :param document: Word 文档对象
    :param text: 段落文本
    """
    try:
        # 将文本拆分成行
        lines = text.split("\n")

        # 添加段落，并在每一行前添加序号和制表符缩进
        for i, line in enumerate(lines, start=1):
            indented_line = f"\t{i}.{line}"
            document.add_paragraph(indented_line, style="BodyText")

    except Exception as e:
        print(f"An error occurred in add_indented_paragraph: {e}")


def word_add_ordered_list(document, items, index_style="number", indentation=1):
    """
    向文档中添加有序列表，并设置缩进。
    :param document: 文档对象。
    :param items: 有序列表项的列表。
    :param index_style: 索引样式，可以是 'number'（数字）、'lower_alpha'（小写字母）、'upper_alpha'（大写字母）、'roman'（罗马数字）之一，默认为 'number'。
    :param indentation: 缩进量，默认为1级0.5英寸。
    """
    try:
        indexes = []  # 存储索引列表
        if index_style == "number":
            indexes = [str(i) for i in range(1, len(items) + 1)]
        elif index_style == "lower_alpha":
            indexes = list(string.ascii_lowercase)[: len(items)]  # 将数字转换为小写字母
        elif index_style == "upper_alpha":
            indexes = [
                string.ascii_lowercase[i].upper() for i in range(len(items))
            ]  # 将数字转换为大写字母
        elif index_style == "roman":
            indexes = [
                common_int_to_roman(i) for i in range(1, len(items) + 1)
            ]  # 使用 to_roman 方法转换数字为罗马数字
        else:
            raise ValueError(
                "Invalid index style. Use 'number', 'low_alpha', 'upper_alpha', or 'roman'."
            )

        for index, item in zip(indexes, items):
            # 如果不是最后一项，在末尾添加分号
            if index != indexes[-1]:
                paragraph = document.add_paragraph(f"{index}、{item}；")
            else:
                # 如果是最后一项，在末尾添加句号
                paragraph = document.add_paragraph(f"{index}、{item}。")
            paragraph.paragraph_format.left_indent = Inches(indentation * 0.5)

    except Exception as e:
        print(f"An error occurred in add_ordered_list: {e}")


def word_add_table_caption(document, caption_text):
    """
    在 Word 文档中添加表序号题注
    :param document: Word 文档对象
    :param caption_text: 题注文本
    """
    try:
        # 获取所有表格
        tables = document.tables

        # 获取表格数量
        table_count = len(tables)

        # 获取最近的标题（章节号）
        section_number = None
        for paragraph in document.paragraphs[::-1]:  # 倒序查找
            if paragraph.style.name.startswith("Heading"):
                section_text = word_get_section_number(paragraph.text)
                if section_text is not None:
                    section_number = section_text.split(".")[0]
                break

        # 创建一个新的段落对象用于存放表序号题注
        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐

        # 创建一个新的表序号题注对象
        caption = OxmlElement("w:fldSimple")
        caption.set(qn("w:instr"), f"SEQ Table \\* ARABIC")
        run = p.add_run()
        run.text = f"\n表 {section_number}-{table_count + 1} {caption_text}"
        run.font.name = "宋体"
        run.font.size = Pt(10.5)
        p._p.append(caption)

    except Exception as e:
        print(f"An error occurred in add_table_caption: {e}")


def word_docxinitial(
        doc,
        text_font_type="宋体",
        text_font_size=12,
        text_font_line_spacing=30,
        text_header="Header",
        text_footer="Footer",
):
    """
    设置文档的全局字体样式、大小和页眉页脚。

    :param doc: Document对象，要设置样式的Word文档对象。
    :param text_font_type: str，可选，正文的字体类型，默认为'宋体'。
    :param text_font_size: int，可选，正文的字体大小，默认为12磅。
    :param text_font_line_spacing: int，可选，正文的行间距，默认为30磅。
    :param text_header: str，可选，页眉内容，默认为'Header'。
    :param text_footer: str，可选，页脚内容，默认为'Footer'。

    :return: Document对象，设置完样式后的Word文档对象。
    """
    # 设置正文字体类型、大小和行间距
    doc.styles["Normal"].font.name = text_font_type
    doc.styles["Normal"].font.size = Pt(text_font_size)
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), text_font_type)
    doc.styles["Normal"]._element.line_spacing = Pt(
        text_font_line_spacing
    )  # 设置行间距为1.5倍

    # 设置页眉
    header = doc.sections[0].header
    pheader = header.paragraphs[0]  # 获取页眉的第一个段落
    ph_header = pheader.add_run(text_header)
    ph_header.font.name = text_font_type  # 设置页眉字体样式
    ph_header._element.rPr.rFonts.set(qn("w:eastAsia"), text_font_type)
    ph_header.font.size = Pt(text_font_size)  # 设置页眉字体大小
    pheader.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 页眉对齐方式设为居中

    # 设置页脚
    footer = doc.sections[0].footer
    pfooter = footer.paragraphs[0]  # 获取页脚的第一个段落
    ph_footer = pfooter.add_run(text_footer)
    ph_footer.font.name = text_font_type  # 设置页脚字体样式
    ph_footer._element.rPr.rFonts.set(qn("w:eastAsia"), text_font_type)
    ph_footer.font.size = Pt(text_font_size)  # 设置页脚字体大小
    pfooter.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 页眉对齐方式设为居中

    # 返回文档对象
    return doc


def word_get_section_number(paragraph_text):
    """
    从段落文本中提取章节号
    """
    try:
        parts = paragraph_text.split()
        for part in parts:
            if part.replace(".", "").isdigit():
                return part
        return None
    except Exception as e:
        print(f"An error occurred in get_section_number: {e}")
        return None


def word_get_table_cell_text(table, row_index, column_index):
    """
    获取表格中指定单元格的文本值。
    :param table: 表格对象。
    :param row_index: 行索引（0开始）。
    :param column_index: 列索引（0开始）。
    :return: 单元格中的文本值。
    """
    try:
        return table.cell(row_index, column_index).text
    except AttributeError as e:
        print(f"AttributeError: {e} occurred in get_table_cell_text")
        return None
    except IndexError as e:
        print(f"IndexError: {e} occurred in get_table_cell_text")
        return None


def word_merge_cells(table, start_row, start_column, end_row, end_column):
    """
    合并表格中指定范围内的单元格。

    参数：
    table: 表格对象
    start_row: 起始行索引
    start_column: 起始列索引
    end_row: 结束行索引
    end_column: 结束列索引
    """
    # 检查起始行和列是否在有效范围内
    if start_row < 0 or start_row >= len(table.rows):
        raise ValueError("Start row index out of range occurred in merge_cells.")
    if start_column < 0 or start_column >= len(table.columns):
        raise ValueError("Start column index out of range occurred in merge_cells.")

    # 检查结束行和列是否在有效范围内
    if end_row < 0 or end_row >= len(table.rows):
        raise ValueError("End row index out of range occurred in merge_cells.")
    if end_column < 0 or end_column >= len(table.columns):
        raise ValueError("End column index out of range occurred in merge_cells.")

    # 合并单元格
    table.cell(start_row, start_column).merge(table.cell(end_row, end_column))


def word_merge_cells_by_column(table, column_index):
    """
    在表格中根据指定列合并相同值的行，并在合并后的单元格中保留唯一值
    :param table: 表格对象
    :param column_index: 要合并的列索引
    """
    try:
        prev_value = None
        merge_start = None

        for row_index, row in enumerate(table.rows):
            current_cell_value = row.cells[column_index].text.strip()

            if current_cell_value != prev_value:
                if merge_start is not None:
                    # Merge cells from merge_start to current row - 1 in the specified column
                    for i in range(merge_start, row_index):
                        # Keep only the value of the first row in the merged cells
                        table.cell(i, column_index).text = prev_value
                        table.cell(i, column_index).merge(
                            table.cell(row_index - 1, column_index)
                        )

                # Update merge_start to the current row
                merge_start = row_index

            # Update prev_value for the next iteration
            prev_value = current_cell_value

        # Merge the last set of cells if needed
        if merge_start is not None and merge_start != len(table.rows):
            for i in range(merge_start, len(table.rows)):
                # Keep only the value of the first row in the merged cells
                table.cell(i, column_index).text = prev_value
                table.cell(i, column_index).merge(
                    table.cell(len(table.rows) - 1, column_index)
                )

    except Exception as e:
        print(f"An error occurred in merge_cells_by_column: {e}")


def word_set_number_to_roman(number):
    """
    将整数转换为罗马数字。
    """
    try:
        roman_numerals = {
            1: "I",
            2: "II",
            3: "III",
            4: "IV",
            5: "V",
            6: "VI",
            7: "VII",
            8: "VIII",
            9: "IX",
            10: "X",
            20: "XX",
            30: "XXX",
            40: "XL",
            50: "L",
            60: "LX",
            70: "LXX",
            80: "LXXX",
            90: "XC",
            100: "C",
            200: "CC",
            300: "CCC",
            400: "CD",
            500: "D",
            600: "DC",
            700: "DCC",
            800: "DCCC",
            900: "CM",
            1000: "M",
        }
        if number in roman_numerals:
            return roman_numerals[number]
        else:
            return "".join(
                [
                    roman_numerals[int(digit) * 10 ** i]
                    for i, digit in enumerate(reversed(str(number)))
                ][::-1]
            )
    except Exception as e:
        print(f"An error occurred in set_number_to_roman: {e}")


def word_set_table_cell_text(table, row_index, column_index, text):
    """
    设置表格中指定单元格的文本内容。

    参数：
    table: 表格对象
    row_index: 行索引（0开始）
    column_index: 列索引（0开始）
    text: 要设置的文本内容
    """
    try:
        table.cell(row_index, column_index).text = text
    except Exception as e:
        print(f"An error occurred in set_table_cell_text: {e}")


def word_set_table_alignment(
        table,
        start_row=None,
        start_column=None,
        end_row=None,
        end_column=None,
        horizontal_alignment=WD_ALIGN_PARAGRAPH.CENTER,
        vertical_alignment=WD_ALIGN_PARAGRAPH.CENTER,
):
    """
    设置表格中指定范围内的单元格的文本对齐方式，默认情况下将所有单元格的文本设置为水平和垂直居中对齐。

    参数：
    table: 表格对象
    start_row: 起始行索引，默认为 None（默认第一行）
    start_column: 起始列索引，默认为 None（默认第一列）
    end_row: 结束行索引，默认为 None（默认最后一行）
    end_column: 结束列索引，默认为 None（默认最后一列）
    horizontal_alignment: 水平对齐方式，默认为水平居中对齐
    vertical_alignment: 垂直对齐方式，默认为垂直居中对齐
    """
    try:
        # 默认情况下，如果未提供范围，则将范围设置为整个表格
        if start_row is None:
            start_row = 0
        if start_column is None:
            start_column = 0
        if end_row is None:
            end_row = len(table.rows) - 1
        if end_column is None:
            end_column = len(table.columns) - 1

        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if (start_row <= i <= end_row and start_column <= j <= end_column) or (
                        start_row > end_row and start_column > end_column
                ):
                    cell.paragraphs[0].alignment = horizontal_alignment
                    cell.vertical_alignment = vertical_alignment

    except Exception as e:
        print(f"An error occurred in set_table_alignment: {e}")


def word_set_table_style(table, font_size=10.5, spacing_before=0):
    """
    设置表格的字体大小和段前间距。
    :param table: 表格对象。
    :param font_size: 字体大小，默认为10.5。
    :param spacing_before: 段前间距，默认为0。
    :return:
    """
    try:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)  # 设置字体大小
                    paragraph.paragraph_format.space_before = Pt(
                        spacing_before
                    )  # 设置段前间距
    except Exception as e:
        print(f"An error occurred in set_table_style: {e}")


def word_setsectionformat(doc):
    """
    在doc结尾处调用进行设置
    通过sections（节）进行设置时，节对应文档中的每一页,每个节在没输入内容之前是不存在的,因此在最后才对每个节逐一进行设置
    :param doc:
    :return:
    """
    for sec in doc.sections:
        # 设置页面边距(左上25毫米，右下15毫米)
        sec.top_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        # 设置纸张大小(A4)
        sec.page_height = Mm(297)
        sec.page_width = Mm(210)
        # 设置页眉页脚距离
        sec.header_distance = Cm(1.5)
        sec.footer_distance = Cm(0.2)
